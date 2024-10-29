import PyPDF2
from docx import Document

# Kullanıcıdan PDF ve Word dosyalarının yollarını isteme
pdf_path = input("PDF dosyasının yolunu girin: ")
word_path = input("Word dosyasının kaydedileceği yol ve dosya adı (örneğin, C:\\Users\\KullaniciAdi\\Downloads\\N_Response_Report.docx): ")

# PDF dosyasını açma
with open(pdf_path, "rb") as pdf_file:
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    document = Document()
    document.add_heading("N-Response Questions and Comments", level=1)

    # Sayfa sayısını ve tüm sayfalardaki metinleri dolaşma
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text = page.extract_text()

        # "N" yanıtı verilen soruları ve "Inspector Observations" bölümlerini bulma
        lines = text.split("\n")
        capturing = False
        question_text = ""

        for line in lines:
            if "Y N" in line or "Y N NA" in line:
                if "N" in line:
                    question_text = line  # Soru metnini kaydedin
                    capturing = True

            elif capturing:
                if "Inspector Observations:" in line:
                    # "N" cevabı ve ilgili açıklamayı yazma
                    question_text += "\n" + line
                    document.add_paragraph(question_text)
                    capturing = False
                    question_text = ""

    # Word belgesini kaydetme
    document.save(word_path)
    print("İşlem tamamlandı! Çıktı Word dosyasına kaydedildi.")
